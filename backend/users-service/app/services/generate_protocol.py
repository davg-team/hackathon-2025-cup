import os
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def generate_certificate(
    full_name: str = None,
    place: str = None,
    output_file: str = None,
    date: str = None,
    signature_name: str = None,
    event: str = None,
):
    # Регистрируем шрифт с футуристичным стилем, поддерживающий кириллицу
    font_path = os.path.join(
        "", "JetBrainsMono-VariableFont_wght.ttf"
    )  # Замените на путь к вашему TTF шрифту
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont("FuturisticFont", font_path))
    else:
        print(f"Шрифт по пути '{font_path}' не найден.")
        return

    # Цвета бренда
    brand_colors = {
        "blue": colors.HexColor("#402FFF"),
        "red": colors.HexColor("#EC1D35"),
        "grey": colors.HexColor("#C8C9CA"),
        "dark": colors.HexColor("#1B1C21"),
        "lightgrey": colors.HexColor("#EDEDED"),
    }

    # Создаем объект Canvas
    c = canvas.Canvas(output_file, pagesize=A4)
    width, height = A4

    # Устанавливаем темный фон
    c.setFillColor(brand_colors["dark"])
    c.rect(0, 0, width, height, stroke=0, fill=1)

    # Добавляем футуристичные элементы дизайна
    c.setStrokeColor(brand_colors["blue"])
    c.setLineWidth(2)
    c.line(2 * cm, height - 4 * cm, width - 2 * cm, height - 4 * cm)
    c.line(2 * cm, 4 * cm, width - 2 * cm, 4 * cm)

    # Добавляем логотип
    logo_path = "logo_light.png"  # Используйте светлый логотип на темном фоне
    if os.path.exists(logo_path):
        logo = ImageReader(logo_path)
        c.drawImage(
            logo,
            width / 2 - 10 * cm,
            height - 3.5 * cm,
            width=20 * cm,
            height=3 * cm,
            mask="auto",
        )
    else:
        print(f"Логотип по пути '{logo_path}' не найден.")

    # Добавляем заголовок
    c.setFont("FuturisticFont", 36)
    c.setFillColor(brand_colors["lightgrey"])
    c.drawCentredString(width / 2, height - 9 * cm, "СЕРТИФИКАТ")

    # Добавляем текст о вручении
    c.setFont("FuturisticFont", 18)
    c.setFillColor(brand_colors["grey"])
    text = f"Настоящим удостоверяется, что\n\n{full_name}\n\nзанял(а) {place}\n"
    print(event)
    if event:
        event_ = ""
        for word in event.split():
            event_ += word + " "
            if len(event.split("\n")[-1]) > 50:
                event_ += "\n"
        event = event_[:]
    print(event)
    text += f"на мероприятии\n{event}" if event else "в нашем мероприятии.\n"
    text_y_start = height - 13 * cm
    line_spacing = 24

    for i, line in enumerate(text.split("\n")):
        c.drawCentredString(width / 2, text_y_start - i * line_spacing, line)

    # Добавляем подпись
    signature_path = "signature_light.png"  # Светлая подпись на темном фоне
    if os.path.exists(signature_path):
        signature = ImageReader(signature_path)
        c.drawImage(
            signature,
            width / 2 - 4 * cm,
            6 * cm,
            width=8 * cm,
            height=2 * cm,
            mask="auto",
        )
    else:
        print(f"Подпись по пути '{signature_path}' не найдена.")

    # Добавляем имя подписывающего
    c.setFont("FuturisticFont", 14)
    c.setFillColor(brand_colors["grey"])
    # c.drawCentredString(
    #     width / 2, 4 * cm, signature_name if signature_name else "Директор мероприятия"
    # )

    # Добавляем дату
    c.setFont("FuturisticFont", 12)
    c.setFillColor(brand_colors["grey"])

    date = date if date else datetime.now().strftime("%d.%m.%Y")
    c.drawRightString(width - 2 * cm, 2 * cm, date if date else "Дата: _____________")

    # Пишеим что это АИС "ФСП"
    c.setFont("FuturisticFont", 12)
    c.setFillColor(brand_colors["grey"])
    c.drawCentredString(width / 2, 2 * cm, 'АИС "ФСП"')

    # Сохранение PDF
    c.showPage()
    c.save()


def generate_protocol(output_file: str, title: str, data: list[dict]):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    title = doc.add_heading(title, level=1)
    title.alignment = 1
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    total_rows = sum(len(team["scores"]) for team in data) + len(data)
    table = doc.add_table(rows=total_rows + 1, cols=8)
    table.style = "Table Grid"
    table.autofit = False
    table.width = Cm(25)

    headers = [
        "Фамилия, имя, отчество\n(при наличии) судьи",
        "Команда (состав)",
        "Работоспособность\nпрототипа (10%)",
        "Соответствие\nфункциональным требованиям (10%)",
        "Презентация (10%)",
        "Технологичность (10%)",
        "Потенциал (10%)",
        "ИТОГО",
    ]

    header_row = table.rows[0]
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = 1

    row_index = 1
    for team in data:
        participants = ", ".join(team["participants"])
        num_judges = len(team["scores"])

        for judge_data in team["scores"]:
            table.cell(row_index, 0).text = judge_data["judge"]
            scores = judge_data["scores"]
            for col_index, score in enumerate(scores, start=2):
                table.cell(row_index, col_index).text = str(score)

            total_score = sum(scores)
            table.cell(row_index, 7).text = str(total_score)
            row_index += 1

        table.cell(row_index, 0).text = "ОБЩЕЕ"
        team_cell = table.cell(row_index - num_judges, 1)
        for merge_row in range(row_index - num_judges, row_index + 1):
            team_cell.merge(table.cell(merge_row, 1))
        team_cell.text = f'{team["team"]}\n({participants})'
        team_cell.paragraphs[0].alignment = 1

        total_team_score = [0] * 5
        for judge_data in team["scores"]:
            for idx, score in enumerate(judge_data["scores"]):
                total_team_score[idx] += score

        for col_index, score in enumerate(total_team_score, start=2):
            table.cell(row_index, col_index).text = str(score)

        table.cell(row_index, 7).text = str(sum(total_team_score))
        row_index += 1

    for row in table.rows:
        cell = row.cells[1]
        cell.width = Cm(5)

    doc.add_paragraph("")
    sign_table = doc.add_table(rows=4, cols=3)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.columns[0].width = Cm(8)
    sign_table.columns[1].width = Cm(6)
    sign_table.columns[2].width = Cm(6)

    sign_table.cell(0, 0).text = "Главный судья соревнований"
    sign_table.cell(0, 1).text = "___________________"
    sign_table.cell(0, 2).text = "___________________"
    sign_table.cell(1, 0).merge(sign_table.cell(1, 0))
    sign_table.cell(1, 1).text = "подпись"
    sign_table.cell(1, 2).text = "инициалы"
    sign_table.cell(2, 0).text = "Главный секретарь соревнований"
    sign_table.cell(2, 1).text = "___________________"
    sign_table.cell(2, 2).text = "___________________"
    sign_table.cell(3, 0).merge(sign_table.cell(3, 0))
    sign_table.cell(3, 1).text = "подпись"
    sign_table.cell(3, 2).text = "инициалы"

    for row in sign_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)

    doc.add_paragraph("\n«___» _______________ 20__ г.")
    doc.add_paragraph("М.П.")
    doc.save(output_file)
